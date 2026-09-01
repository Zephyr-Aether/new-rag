"""知识库文档导入（文件上传解析）：PDF/DOC/DOCX/Markdown/TXT → 文本；CSV/Excel → FAQ 问答对。

按文件魔数嗅探真实格式（%PDF / OLE2 / ZIP），不依赖扩展名——docx 与 xlsx 同为 ZIP、
doc 与 xls 同为 OLE2，靠容器内部目录区分。docx 优先用 python-docx（缺库或解析失败时回退
zip 内联抽取），.doc 用 olefile 读 WordDocument 流的 FIB fcMin/fcMac 定位 UTF-16LE 正文。
FAQ 文件每行（问题, 答案）转为 `## 问：…\n答：…` 小节，配合现有 Markdown 分块器按小节切块。
上传文本先经 clean_document 清洗（去噪/规范化），再入库。
"""

import csv
import html
import io
import re
import struct
import zipfile

# 常见噪音：独立页码行（- 1 - / — 3 — / 第 1 页 / 页 3 / 纯数字）、分隔线、PDF 页眉页脚
_PAGE_NO_RE = re.compile(
    r"^\s*(?:[-—–_]{1,}\s*\d+\s*[-—–_]{1,}|第\s*\d+\s*页|页\s*\d+|\d+\s*/\s*\d+|\d{1,4})\s*$"
)
_CONTROL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
_ZERO_WIDTH_RE = re.compile(r"[​-‍﻿]")
_WS_LINE_RE = re.compile(r"[ \t]+$")

# 文件魔数
_OLE2_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"
_ZIP_MAGICS = (b"PK\x03\x04", b"PK\x05\x06", b"PK\x07\x08")


def clean_document(text: str) -> str:
    """文档清洗：统一换行、去控制字符/零宽、去独立页码行、压缩空行与行尾空格。"""
    text = _ZERO_WIDTH_RE.sub("", text or "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _CONTROL_RE.sub("", text)
    lines = []
    for line in text.split("\n"):
        line = _WS_LINE_RE.sub("", line).strip()
        if not line:
            lines.append("")
            continue
        if _PAGE_NO_RE.match(line):
            continue
        lines.append(line)
    # 压缩连续空行（最多 1 个）
    out: list[str] = []
    prev_blank = False
    for line in lines:
        if not line:
            if not prev_blank:
                out.append("")
            prev_blank = True
        else:
            out.append(line)
            prev_blank = False
    return "\n".join(out).strip()


class ImportError_(Exception):
    pass


def _decode(data: bytes) -> str:
    for enc in ("utf-8", "gb18030"):
        try:
            return data.decode(enc)
        except (UnicodeDecodeError, UnicodeError):
            continue
    return data.decode("utf-8", errors="replace")


def _parse_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:  # noqa: BLE001 单页失败跳过
            text = ""
        if text.strip():
            pages.append(text)
    return "\n\n".join(pages)


def _parse_docx(data: bytes) -> str:
    """DOCX → 文本：优先 python-docx（段落+表格），缺库/解析失败时回退 zip 内联抽取。"""
    try:
        text = _parse_docx_pydocx(data)
    except Exception:  # noqa: BLE001 交给 zip 兜底
        text = ""
    if text.strip():
        return text
    return _parse_docx_zip(data)


def _parse_docx_pydocx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        return ""
    doc = Document(io.BytesIO(data))
    parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            # 合并单元格会重复；dict.fromkeys 保序去重
            cells = list(dict.fromkeys(c.text.strip() for c in row.cells))
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def _parse_docx_zip(data: bytes) -> str:
    """零依赖兜底：抽 word/document.xml，剥标签并保留段落/表格/换行结构。"""
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        xml = z.read("word/document.xml").decode("utf-8", errors="replace")
    xml = re.sub(r"<w:tab[^>]*/>", "\t", xml)
    xml = re.sub(r"<w:br[^>]*/>", "\n", xml)
    xml = re.sub(r"</w:p>|</w:tr>", "\n", xml)
    xml = re.sub(r"</w:tc>", "\t", xml)
    text = re.sub(r"<[^>]+>", "", xml)
    return html.unescape(text)


def _garbage_ratio(text: str) -> float:
    if not text:
        return 1.0
    bad = text.count("\ufffd")
    bad += sum(1 for ch in text if ord(ch) < 32 and ch not in "\t\n\r")
    return bad / len(text)


def _parse_doc(data: bytes) -> str:
    """旧版 Word .doc（OLE2）→ 文本：读 WordDocument 流，按 FIB fcMin/fcMac 取 UTF-16LE 正文。"""
    import olefile

    ole = olefile.OleFileIO(io.BytesIO(data))
    try:
        if not ole.exists("WordDocument"):
            raise ImportError_("不是有效的 Word .doc 文件（缺少 WordDocument 流）")
        stream = bytearray(ole.openstream("WordDocument").read())
        if len(stream) < 0x20:
            raise ImportError_("WordDocument 流过短，无法解析")
        flags = struct.unpack_from("<H", stream, 0x0A)[0]
        fc_min = struct.unpack_from("<I", stream, 0x18)[0]
        fc_mac = struct.unpack_from("<I", stream, 0x1C)[0]
        # fComplex=1 时正文在表流（0Table/1Table），fcMin/fcMac 指向表流
        if flags & 0x0004:
            table_name = "1Table" if flags & 0x0080 else "0Table"
            if ole.exists(table_name):
                stream = bytearray(ole.openstream(table_name).read())
        if fc_min >= fc_mac or fc_mac > len(stream):
            raise ImportError_("无法定位 .doc 正文（文件结构异常）")
        raw = bytes(stream[fc_min:fc_mac])
        if raw.endswith(b"\x0d\x00"):  # 段落符 \r 以 UTF-16LE 存储
            raw = raw[:-2]
        text = raw.decode("utf-16-le", errors="replace")
        if _garbage_ratio(text) > 0.05:
            text = raw.decode("gb18030", errors="replace")  # 老式 8-bit 码页（中文 GBK）
        if _garbage_ratio(text) > 0.05:
            raise ImportError_("无法从 .doc 提取文本（复杂版式请转存为 .docx）")
        return text
    finally:
        ole.close()


def _rows_to_faq_markdown(rows: list[list[str]]) -> str:
    blocks = []
    for row in rows:
        if len(row) < 2 or not (row[0] or "").strip():
            continue
        q, a = (row[0] or "").strip(), (row[1] or "").strip()
        blocks.append(f"## 问：{q}\n答：{a}")
    return "\n\n".join(blocks)


def _parse_csv(data: bytes) -> str:
    text = _decode(data)
    rows = list(csv.reader(io.StringIO(text)))
    return _rows_to_faq_markdown(rows)


def _parse_excel(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    ws = wb.active
    rows = [[(c if c is not None else "") for c in row] for row in ws.iter_rows(values_only=True)]
    return _rows_to_faq_markdown(rows)


def _zip_names(data: bytes) -> list[str]:
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            return z.namelist()
    except Exception:  # noqa: BLE001
        return []


def _ole_streams(data: bytes) -> set[str]:
    try:
        import olefile

        ole = olefile.OleFileIO(io.BytesIO(data))
        try:
            return {p[-1] for p in ole.listdir() if p}
        finally:
            ole.close()
    except Exception:  # noqa: BLE001
        return set()


def _detect_format(filename: str, data: bytes) -> str:
    """嗅探真实格式（优先魔数），返回 pdf/docx/doc/xlsx/xls/csv/text/zip/ole。"""
    if data.startswith(b"%PDF"):
        return "pdf"
    if data.startswith(_ZIP_MAGICS):
        names = _zip_names(data)
        if any(n.startswith("word/") for n in names):
            return "docx"
        if any(n.startswith("xl/") for n in names):
            return "xlsx"
        return "zip"
    if data.startswith(_OLE2_MAGIC):
        streams = _ole_streams(data)
        if "WordDocument" in streams:
            return "doc"
        if "Workbook" in streams or "Book" in streams:
            return "xls"
        return "ole"
    name = (filename or "").lower()
    if name.endswith(".csv"):
        return "csv"
    return "text"


def parse_document(filename: str, data: bytes) -> str:
    """按真实格式解析上传文件为 Markdown 文本，并做文档清洗。返回空串表示无可提取内容。

    文件大小上限由上层接口层校验（settings.max_upload_bytes / upload_manager.max_bytes），
    这里不做重复限制，避免与 50MB 配置冲突。
    """
    fmt = _detect_format(filename, data)
    try:
        if fmt == "pdf":
            text = _parse_pdf(data)
        elif fmt == "docx":
            text = _parse_docx(data)
        elif fmt == "doc":
            text = _parse_doc(data)
        elif fmt == "xlsx":
            text = _parse_excel(data)
        elif fmt == "xls":
            raise ImportError_("不支持旧版 .xls 二进制格式，请另存为 .xlsx 后重试")
        elif fmt == "csv":
            text = _parse_csv(data)
        elif fmt == "text":
            text = _decode(data)
        else:
            raise ImportError_("暂不支持的文件类型，请上传 PDF / DOC / DOCX / XLSX / CSV / TXT / MD")
    except ImportError_:
        raise
    except Exception as exc:  # noqa: BLE001
        raise ImportError_(f"parse failed: {exc}") from exc
    return clean_document(text)
