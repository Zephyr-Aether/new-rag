"""知识库导入解析（importers）：格式嗅探 + PDF/DOCX/DOC/CSV/XLSX/TXT 解析。

PDF 用手工构造的最小合法文件；DOCX 用 python-docx 生成；.doc 用最小 OLE2 容器手工构造
（扇区 0=FAT / 1=目录 / 2+=数据，正文 UTF-16LE，FIB 提供 fcMin/fcMac）。
"""

import io
import struct

import pytest

from app.knowledge.importers import (
    ImportError_,
    _detect_format,
    _parse_docx_zip,
    parse_document,
)

DOCX_TEXT = "退款到账时间是 3 个工作日"
DOC_TEXT = "这是 .doc 正文：退款 3-5 个工作日到账"


# ---------- 样本文件构造 ----------


def _build_pdf(text: str) -> bytes:
    """最小合法 PDF：1 页 + Helvetica 文本 + 正确 xref。"""
    content = f"BT /F1 24 Tf 100 700 Td ({text}) Tj ET".encode("ascii")
    objs = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R "
        b"/Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length %d >>\nstream\n%s\nendstream" % (len(content), content),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for i, obj in enumerate(objs, 1):
        offsets.append(len(out))
        out += b"%d 0 obj\n" % i + obj + b"\nendobj\n"
    xref = len(out)
    out += b"xref\n0 6\n0000000000 65535 f \n"
    for off in offsets[1:]:
        out += b"%010d 00000 n \n" % off
    out += b"trailer\n<< /Size 6 /Root 1 0 R >>\nstartxref\n%d\n%%%%EOF\n" % xref
    return bytes(out)


def _build_docx() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_paragraph(DOCX_TEXT)
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "问题"
    table.cell(0, 1).text = "答案"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_xlsx() -> bytes:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.append(["问题", "答案"])
    ws.append(["退款到账时间", "3 个工作日"])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _build_ole(stream_name: str, payload: bytes) -> bytes:
    """最小 OLE2 容器：扇区 0=FAT / 1=目录 / 2+=数据链。payload 必须 >4096 走常规 FAT。"""
    payload = payload.ljust((len(payload) + 511) // 512 * 512, b"\x00")
    size = len(payload)
    n_data = size // 512
    fat = [0xFFFFFFFF] * 128
    fat[0] = 0xFFFFFFFE  # 扇区 0 = FAT 本身
    fat[1] = 0xFFFFFFFD  # 扇区 1 = 目录链结束
    for i in range(n_data - 1):
        fat[2 + i] = 3 + i
    fat[2 + n_data - 1] = 0xFFFFFFFD

    def entry(name: str, etype: int, child: int, start: int, esize: int) -> bytes:
        e = bytearray(128)
        nb = name.encode("utf-16-le") + b"\x00\x00"
        e[0 : len(nb)] = nb
        struct.pack_into("<H", e, 64, len(name) * 2 + 2)
        e[66] = etype  # 5=root, 2=stream
        e[67] = 1
        struct.pack_into("<i", e, 68, -1)
        struct.pack_into("<i", e, 72, -1)
        struct.pack_into("<i", e, 76, child)
        struct.pack_into("<I", e, 116, start)
        struct.pack_into("<Q", e, 120, esize)
        return bytes(e)

    dir_sector = (
        entry("Root Entry", 5, 1, 0xFFFFFFFE, 0) + entry(stream_name, 2, -1, 2, size) + b"\x00" * 256
    )
    hdr = bytearray(512)
    hdr[0:8] = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"
    struct.pack_into("<H", hdr, 0x18, 0x003E)
    struct.pack_into("<H", hdr, 0x1A, 0x0003)
    struct.pack_into("<H", hdr, 0x1C, 0xFFFE)
    struct.pack_into("<H", hdr, 0x1E, 9)
    struct.pack_into("<H", hdr, 0x20, 6)
    struct.pack_into("<I", hdr, 0x28, 1)
    struct.pack_into("<I", hdr, 0x2C, 1)
    struct.pack_into("<I", hdr, 0x30, 1)  # 目录扇区 1
    struct.pack_into("<I", hdr, 0x38, 4096)
    struct.pack_into("<I", hdr, 0x3C, 0xFFFFFFFE)
    struct.pack_into("<I", hdr, 0x40, 0)
    struct.pack_into("<I", hdr, 0x44, 0xFFFFFFFE)
    struct.pack_into("<I", hdr, 0x48, 0)
    struct.pack_into("<I", hdr, 0x4C, 0)  # DIFAT[0] → FAT 扇区 0
    for i in range(1, 109):
        struct.pack_into("<I", hdr, 0x4C + 4 * i, 0xFFFFFFFF)
    return bytes(hdr) + b"".join(struct.pack("<I", x) for x in fat) + dir_sector + payload


def _build_doc(text: str) -> bytes:
    """最小可解析 Word 97 .doc：OLE2 容器 + WordDocument 流（FIB + UTF-16LE 正文）。"""
    body = text.encode("utf-16-le")
    fib = bytearray(0x200)
    struct.pack_into("<H", fib, 0x00, 0xA5EC)  # wIdent
    struct.pack_into("<H", fib, 0x02, 0x00C1)  # nFib (Word 97)
    struct.pack_into("<I", fib, 0x18, 0x200)  # fcMin
    struct.pack_into("<I", fib, 0x1C, 0x200 + len(body))  # fcMac
    # 补到 >4096 走常规 FAT（构造器未实现 mini stream），fcMac 之后是填充，不影响正文切片
    return _build_ole("WordDocument", (bytes(fib) + body).ljust(4608, b"\x00"))


# ---------- 格式嗅探 ----------


class TestDetectFormat:
    def test_pdf_by_magic(self):
        assert _detect_format("report.pdf", _build_pdf("Hi")) == "pdf"
        # 魔数优先于扩展名
        assert _detect_format("note.txt", _build_pdf("Hi")) == "pdf"

    def test_docx_vs_xlsx_both_zip(self):
        assert _detect_format("a.docx", _build_docx()) == "docx"
        assert _detect_format("b.docx", _build_xlsx()) == "xlsx"  # 内容才是真相

    def test_doc_by_ole_magic(self):
        doc = _build_doc(DOC_TEXT)
        assert _detect_format("note.doc", doc) == "doc"
        # 扩展名写 docx 也不影响魔数判定
        assert _detect_format("note.docx", doc) == "doc"

    def test_csv_by_extension(self):
        assert _detect_format("faq.csv", b"q,a\n") == "csv"
        assert _detect_format("plain.txt", b"hello") == "text"


# ---------- 解析 ----------


class TestParseDocument:
    def test_pdf(self):
        text = parse_document("report.pdf", _build_pdf("Hello World"))
        assert "Hello World" in text

    def test_docx_paragraphs_and_table(self):
        text = parse_document("a.docx", _build_docx())
        assert DOCX_TEXT in text
        assert "问题" in text and "答案" in text

    def test_docx_zip_fallback(self):
        """python-docx 不可用时走 zip 兜底，仍能抽到正文。"""
        text = _parse_docx_zip(_build_docx())
        assert DOCX_TEXT in text

    def test_doc(self):
        text = parse_document("legacy.doc", _build_doc(DOC_TEXT))
        assert DOC_TEXT in text

    def test_doc_extension_ignored_magic_wins(self):
        text = parse_document("trick.docx", _build_doc(DOC_TEXT))
        assert DOC_TEXT in text

    def test_csv_faq(self):
        text = parse_document("faq.csv", "问题一,答案一\n问题二,答案二\n".encode())
        assert "## 问：问题一" in text and "答：答案一" in text

    def test_xlsx_faq(self):
        text = parse_document("faq.xlsx", _build_xlsx())
        assert "## 问：退款到账时间" in text and "答：3 个工作日" in text

    def test_txt_chinese_gb18030(self):
        raw = "中文测试".encode("gb18030")
        assert parse_document("note.txt", raw) == "中文测试"

    def test_xls_unsupported_clear_error(self):
        xls = _build_ole("Workbook", b"\x00" * 4608)
        with pytest.raises(ImportError_, match="xlsx"):
            parse_document("old.xls", xls)

    def test_unknown_zip_rejected(self):
        import zipfile

        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as z:
            z.writestr("random.txt", "not a document container")
        with pytest.raises(ImportError_, match="暂不支持"):
            parse_document("thing.zip", buf.getvalue())

    def test_empty_text_returns_empty(self):
        # 无内容 → 返回空串（转成友好报错在 api 层做）
        assert parse_document("empty.txt", b"") == ""
